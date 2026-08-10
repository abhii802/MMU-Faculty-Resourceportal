#!/usr/bin/env python3
"""
CS Notes Generator
==================
A modern, beautifully formatted notes generator for Computer Science topics.
Generates DOCX with diagrams, real-world examples, and beginner-friendly language.

Usage:
    python cs_notes_generator.py

Requirements:
    pip install python-docx Pillow matplotlib requests

Optional (for PDF conversion on Windows):
    pip install docx2pdf

For Linux/Mac PDF conversion, install LibreOffice and run:
    libreoffice --headless --convert-to pdf output.docx
"""

import os
import sys
import json
import textwrap
import tempfile
from datetime import datetime
from io import BytesIO
from typing import List, Dict, Optional, Tuple

# Try importing optional dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not installed. Run: pip install python-docx")

try:
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    print("WARNING: matplotlib not installed. Run: pip install matplotlib")

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False


# ============================================================
# CONFIGURATION
# ============================================================
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
USE_AI = bool(OPENAI_API_KEY or ANTHROPIC_API_KEY)

# Modern color scheme
COLORS = {
    "primary": "2D5AF5",
    "secondary": "7B61FF",
    "accent": "00C9A7",
    "dark": "1A1A2E",
    "light": "F8F9FA",
    "warning": "FF6B6B",
    "success": "51CF66",
    "text": "2D3436",
    "muted": "636E72",
}

# ============================================================
# KNOWLEDGE BASE (Fallback when no API key available)
# ============================================================
KNOWLEDGE_BASE = {
    "arrays": {
        "title": "Arrays & Lists",
        "overview": "Arrays are like a row of lockers in a school hallway. Each locker has a number (index), and you can store one item in each locker. All lockers are the same size and arranged one after another in memory.",
        "key_concepts": [
            ("Index", "The locker number. In programming, we start counting from 0, not 1. So the first item is at index 0, second at index 1, and so on."),
            ("Element", "The actual data stored in each locker — could be a number, text, or even another array."),
            ("Length/Size", "Total number of lockers in the row. Once created, the size of a basic array is fixed."),
            ("Contiguous Memory", "All lockers sit next to each other in the computer's memory, making it fast to jump to any locker."),
        ],
        "real_world_examples": [
            ("Movie Theater Seats", "Think of a movie theater with numbered seats (A1, A2, A3...). The seat number is the index, and the person sitting there is the element. You can instantly find seat A10 because you know exactly where it is."),
            ("Egg Carton", "An egg carton is a 2D array! It has rows and columns. Egg at position [0][2] means first row, third column."),
            ("Contacts List", "Your phone's contact list is a dynamic array. You can scroll to contact #50 instantly because the phone knows exactly where each entry is stored."),
        ],
        "pros_cons": [
            ("Fast Access", "O(1) — Instantly access any element by index, like knowing exactly which locker to open."),
            ("Memory Efficient", "No extra space needed for links or pointers — just pure data."),
            ("Fixed Size", "Static arrays cannot grow. Like buying a 10-locker cabinet — if you need 11, you must buy a whole new one."),
            ("Costly Insertion", "Inserting in the middle requires shifting all later elements, like asking everyone in a queue to move one step back."),
        ],
        "code_example": """# Creating an array in Python (list)
scores = [85, 92, 78, 90, 88]

# Accessing elements (O(1))
first_score = scores[0]   # 85
last_score = scores[-1]   # 88

# Updating (O(1))
scores[2] = 95  # Changed 78 to 95

# Inserting in middle (O(n))
scores.insert(2, 100)  # Everyone after index 2 shifts right

# Traversing (O(n))
for i, score in enumerate(scores):
    print(f"Student {i}: {score}")""",
        "diagram_type": "array_structure"
    },

    "linked list": {
        "title": "Linked Lists",
        "overview": "A Linked List is like a treasure hunt. Each clue (node) contains some treasure (data) and a hint (pointer) to where the next clue is hidden. Unlike arrays, the clues don't have to be in order — they can be scattered anywhere!",
        "key_concepts": [
            ("Node", "Each clue in the treasure hunt. Contains: (1) the treasure (data), and (2) the hint to next clue (pointer)."),
            ("Head", "The starting point — the first clue that kicks off the treasure hunt."),
            ("Pointer/Reference", "The hint written on each clue telling you where to find the next one."),
            ("Tail", "The last clue that says 'The End' — its pointer is NULL/None."),
        ],
        "real_world_examples": [
            ("Treasure Hunt", "Each clue paper says 'Look under the bench' — you go there and find the next clue. The clues aren't stored in order; they're scattered around the park."),
            ("Train Cars", "A train is a linked list! Each car (node) is connected to the next via a coupling (pointer). You can add/remove cars without moving the entire train."),
            ("Music Playlist", "Spotify's 'Next Song' is a pointer. Each song knows which song comes next. You can rearrange the playlist by just changing the 'next' pointers, without moving the actual song files."),
        ],
        "pros_cons": [
            ("Dynamic Size", "Grows and shrinks easily. Like adding a new train car — just connect the coupling."),
            ("Easy Insert/Delete", "O(1) if you know the position. Just change the pointer/hint — no shifting needed!"),
            ("No Random Access", "O(n) to find element #100. You must follow all hints from the start, like going through 100 clues."),
            ("Extra Memory", "Each node needs space for the pointer, like each clue needing extra paper for the hint."),
        ],
        "code_example": """class Node:
    def __init__(self, data):
        self.data = data
        self.next = None  # Pointer to next clue

class LinkedList:
    def __init__(self):
        self.head = None

    def append(self, data):
        new_node = Node(data)
        if not self.head:
            self.head = new_node
            return
        last = self.head
        while last.next:
            last = last.next
        last.next = new_node

    def display(self):
        current = self.head
        while current:
            print(current.data, end=" -> ")
            current = current.next
        print("None")

# Usage
ll = LinkedList()
ll.append("A")
ll.append("B")
ll.append("C")
ll.display()  # A -> B -> C -> None""",
        "diagram_type": "linked_list"
    },

    "stack": {
        "title": "Stacks (LIFO)",
        "overview": "A Stack is exactly like a stack of plates in a cafeteria. You can only add (push) a new plate on top, and you can only take (pop) the top plate off. The last plate you put on is the first one you take off — Last In, First Out (LIFO).",
        "key_concepts": [
            ("Push", "Adding a new plate on top of the stack. The new plate becomes the top."),
            ("Pop", "Removing the top plate. You cannot grab a plate from the middle without toppling the stack!"),
            ("Peek/Top", "Looking at the top plate without removing it."),
            ("Underflow", "Trying to pop from an empty stack — like trying to take a plate when there are none."),
            ("Overflow", "Trying to push when the stack is full (in fixed-size implementations)."),
        ],
        "real_world_examples": [
            ("Stack of Plates", "The classic example! You wash a plate and put it on top. When someone needs a plate, they take the top one. The last washed plate is used first."),
            ("Browser Back Button", "Every page you visit is 'pushed' onto a stack. Clicking 'Back' pops the current page off, revealing the previous one."),
            ("Undo in Word/Photoshop", "Every action you take is pushed onto a stack. Ctrl+Z pops the last action off, undoing it. Ctrl+Y (redo) pushes it back on."),
            ("Call Stack", "When function A calls function B, which calls function C — they stack up. When C finishes, it pops off, returning to B, then B pops, returning to A."),
        ],
        "pros_cons": [
            ("Simple & Fast", "Push and Pop are O(1) — instant, like placing or grabbing the top plate."),
            ("Memory Efficient", "Only needs to track the top element."),
            ("Limited Access", "Cannot access middle elements directly. You must pop everything above it first."),
            ("Fixed Size Risk", "Array-based stacks can overflow if not sized properly."),
        ],
        "code_example": """class Stack:
    def __init__(self):
        self.items = []

    def push(self, item):
        self.items.append(item)  # Add to top

    def pop(self):
        if not self.is_empty():
            return self.items.pop()  # Remove from top
        return None

    def peek(self):
        if not self.is_empty():
            return self.items[-1]  # Look at top
        return None

    def is_empty(self):
        return len(self.items) == 0

# Real-world: Undo feature
history = Stack()
history.push("Typed 'Hello'")
history.push("Bold text")
history.push("Added image")

print(history.pop())  # Undo: Added image
print(history.peek()) # Current state: Bold text""",
        "diagram_type": "stack_operations"
    },

    "queue": {
        "title": "Queues (FIFO)",
        "overview": "A Queue is like a line at a movie ticket counter. The first person to arrive gets served first — First In, First Out (FIFO). New people join at the back (enqueue), and people leave from the front (dequeue).",
        "key_concepts": [
            ("Enqueue", "Joining the line at the back. Also called 'offer' or 'add'."),
            ("Dequeue", "Leaving the line from the front after being served. Also called 'poll' or 'remove'."),
            ("Front/Head", "The person currently being served — the first in line."),
            ("Rear/Tail", "The last person in line — the most recent arrival."),
        ],
        "real_world_examples": [
            ("Movie Ticket Line", "People arrive and stand at the back. The ticket seller serves the person at the front. No cutting in line allowed!"),
            ("Printer Queue", "When 5 people send print jobs, the printer handles them in order. First document sent prints first — even if it's 1 page and the next is 100 pages."),
            ("YouTube Buffering", "Video data arrives in a queue. The first chunk downloaded plays first while new chunks join the back of the queue."),
            ("Customer Service Hotline", "Calls are placed in a queue. 'You are caller #5 in line' — you must wait for the 4 people ahead of you."),
        ],
        "pros_cons": [
            ("Fair Ordering", "Guaranteed fairness — everyone served in arrival order."),
            ("Fast Operations", "Enqueue and Dequeue are O(1) with proper implementation."),
            ("Inefficient Array Implementation", "Simple array queues waste space because removing from front requires shifting everyone left."),
            ("Priority Ignored", "A small urgent job must wait behind a huge slow job."),
        ],
        "code_example": """from collections import deque

class Queue:
    def __init__(self):
        self.items = deque()

    def enqueue(self, item):
        self.items.append(item)  # Join at back

    def dequeue(self):
        if not self.is_empty():
            return self.items.popleft()  # Serve from front
        return None

    def front(self):
        return self.items[0] if self.items else None

    def is_empty(self):
        return len(self.items) == 0

# Real-world: Printer queue
printer = Queue()
printer.enqueue("Resume.pdf")
printer.enqueue("Report.docx")
printer.enqueue("Photo.jpg")

print(printer.dequeue())  # Resume.pdf prints first
print(printer.front())    # Next: Report.docx""",
        "diagram_type": "queue_structure"
    },

    "binary search": {
        "title": "Binary Search",
        "overview": "Binary Search is like finding a word in a dictionary. Instead of checking every page from start to end, you open the middle. If your word comes before the middle, you search the left half; if after, the right half. You keep cutting the book in half until you find the word!",
        "key_concepts": [
            ("Sorted Data Required", "The dictionary must be alphabetically ordered. Binary search ONLY works on sorted arrays/lists."),
            ("Midpoint", "Always check the middle element first. This is the 'magic' that makes it fast."),
            ("Divide & Conquer", "Cut the problem in half each time. One half is instantly eliminated!"),
            ("Time Complexity", "O(log n) — Finding one item in 1 billion sorted items takes only 30 checks!"),
        ],
        "real_world_examples": [
            ("Dictionary Lookup", "Finding 'elephant' in a dictionary: Open middle (M), E is before M, so ignore the entire second half. Open middle of first half (G), E is before G... keep halving!"),
            ("Guessing Game", "Think of a number between 1 and 100. Guess 50. If too high, guess 25. If too low, guess 37. You'll find it in 7 guesses max!"),
            ("Git Bisect", "When a bug appeared between version 1.0 and 100.0, Git tests version 50. If buggy, bug is in first half. If not, in second half. Finds the exact commit in log(n) steps."),
            ("IP Address Routing", "Routers use binary search on sorted routing tables to find where to send your packet — millions of routes checked in microseconds."),
        ],
        "pros_cons": [
            ("Blazing Fast", "O(log n) — 1 million items need only 20 checks. 1 billion need only 30!"),
            ("Simple Logic", "Just compare, then go left or right."),
            ("Must Be Sorted", "Won't work on unsorted data. Sorting first costs O(n log n)."),
            ("Random Access Needed", "Requires arrays (O(1) access). Linked lists are too slow for binary search."),
        ],
        "code_example": """def binary_search(arr, target):
    left, right = 0, len(arr) - 1

    while left <= right:
        mid = (left + right) // 2

        if arr[mid] == target:
            return mid  # Found it!
        elif arr[mid] < target:
            left = mid + 1  # Search right half
        else:
            right = mid - 1  # Search left half

    return -1  # Not found

# Finding 'Mango' in sorted fruit list
fruits = ["Apple", "Banana", "Cherry", "Grape", "Mango", 
          "Orange", "Peach", "Strawberry", "Watermelon"]

index = binary_search(fruits, "Mango")
print(f"Mango is at index: {index}")  # Index 4

# Only 3 checks needed instead of checking all 9!""",
        "diagram_type": "binary_search_tree"
    },

    "recursion": {
        "title": "Recursion",
        "overview": "Recursion is like a set of Russian nesting dolls. Each doll contains a smaller version of itself inside. To see the smallest doll, you must keep opening dolls until you find one that can't be opened (base case). Then you close them all back up, returning to the original.",
        "key_concepts": [
            ("Base Case", "The smallest doll that doesn't open — the condition that stops the recursion. WITHOUT THIS, you get infinite recursion!"),
            ("Recursive Case", "The doll opening to reveal a smaller doll — the function calling itself with a smaller problem."),
            ("Call Stack", "Each opened doll is placed on a table (stack). You can't close the outer doll until the inner one is closed."),
            ("Divide & Conquer", "Break a big problem into smaller identical problems until they're trivial to solve."),
        ],
        "real_world_examples": [
            ("Russian Matryoshka Dolls", "Open a doll -> find a smaller doll -> open it -> ... -> find the tiny solid doll (base case). Then close them all in reverse order."),
            ("Factorial", "5! = 5 x 4! = 5 x 4 x 3! = ... = 5 x 4 x 3 x 2 x 1. Each factorial asks a smaller factorial for help."),
            ("Tree of Folders", "To count all files on your computer: Count files in this folder + recursively count in each subfolder. The base case is an empty folder (0 files)."),
            ("Dream within a Dream (Inception)", "Cobb goes deeper into dreams. Each level is a smaller version of reality. The 'kick' (base case) wakes them up level by level."),
        ],
        "pros_cons": [
            ("Elegant Code", "Complex problems become simple, readable solutions."),
            ("Natural for Trees", "Perfect for file systems, DOM traversal, and organizational hierarchies."),
            ("Stack Overflow Risk", "Too many nested calls crash the program — like having infinite dolls."),
            ("Memory Heavy", "Each call uses stack memory. Iterative solutions are often more memory-efficient."),
        ],
        "code_example": """def factorial(n):
    # Base case: smallest doll
    if n <= 1:
        return 1

    # Recursive case: ask smaller factorial for help
    return n * factorial(n - 1)

print(factorial(5))  # 120
# 5! = 5 x 4! = 5 x 4 x 3! = 5 x 4 x 3 x 2! = 5 x 4 x 3 x 2 x 1! = 120


def count_files(folder):
    # Base case: empty folder
    if not folder.subfolders:
        return len(folder.files)

    # Recursive case: count here + count in all children
    total = len(folder.files)
    for sub in folder.subfolders:
        total += count_files(sub)
    return total""",
        "diagram_type": "recursion_tree"
    },

    "object oriented programming": {
        "title": "Object-Oriented Programming (OOP)",
        "overview": "OOP is like a car factory. Instead of building each car from scratch with loose parts, you create a blueprint (Class) that defines what every car should have (attributes) and what it can do (methods). Then you stamp out identical cars (Objects) from that blueprint.",
        "key_concepts": [
            ("Class", "The blueprint or cookie cutter. Defines WHAT exists (color, speed) and WHAT can be done (drive, brake)."),
            ("Object/Instance", "An actual car built from the blueprint. My red Toyota is an object; your blue Honda is another object — both from the 'Car' class."),
            ("Encapsulation", "The car hides its engine complexity. You just turn the key — you don't need to know how combustion works. Data is protected inside the object."),
            ("Inheritance", "A SportsCar blueprint inherits from the Car blueprint and adds 'turbo boost'. It gets all Car features for free, plus new ones."),
            ("Polymorphism", "The 'Start' button works on any vehicle — Car, Motorcycle, or Boat — but each does it differently. Same command, different behavior."),
            ("Abstraction", "The steering wheel is an abstraction. You don't see the rack-and-pinion gears; you just turn the wheel and the car responds."),
        ],
        "real_world_examples": [
            ("Car Factory", "One blueprint (Class) -> thousands of cars (Objects). Change the blueprint, and all future cars change. But existing cars stay as they were built."),
            ("Animal Kingdom", "Class 'Animal' has eat() and sleep(). Class 'Dog' INHERITS from Animal and adds bark(). Class 'Cat' also inherits but adds meow()."),
            ("Restaurant Menu", "The menu (Class) lists dishes and prices. Your order (Object) is a specific instance: 'Cheeseburger with extra cheese, no onions'."),
            ("Bank Account", "Class 'Account' has balance (attribute) and deposit/withdraw (methods). Encapsulation prevents direct balance tampering — you MUST use methods."),
        ],
        "pros_cons": [
            ("Reusability", "Write once (class), use infinitely (objects)."),
            ("Organized Code", "Real-world modeling makes code intuitive and maintainable."),
            ("Steep Learning Curve", "Abstraction, inheritance, and polymorphism can confuse beginners."),
            ("Overhead", "Simple programs may become unnecessarily complex with OOP."),
        ],
        "code_example": """class Car:
    def __init__(self, brand, color):
        self.brand = brand      # Attribute
        self.color = color
        self.speed = 0

    def accelerate(self):       # Method
        self.speed += 10
        print(f"{self.color} {self.brand} speeding up to {self.speed} km/h")

    def brake(self):
        self.speed = 0
        print("Stopped!")

# Creating objects from the blueprint
my_car = Car("Toyota", "Red")
friend_car = Car("Honda", "Blue")

my_car.accelerate()      # Red Toyota speeding up...
friend_car.accelerate()  # Blue Honda speeding up...


class ElectricCar(Car):   # Inheritance
    def __init__(self, brand, color, battery):
        super().__init__(brand, color)
        self.battery = battery

    def charge(self):
        print(f"Charging {self.battery}kWh battery...")

tesla = ElectricCar("Tesla", "White", 75)
tesla.accelerate()  # Inherited from Car!
tesla.charge()      # New method""",
        "diagram_type": "oop_hierarchy"
    },

    "database": {
        "title": "Databases & SQL",
        "overview": "A Database is like a super-organized digital filing cabinet. Instead of throwing papers in a pile, you use labeled drawers (Tables), sorted folders (Rows), and specific fields (Columns). SQL is the language you use to ask the cabinet for exactly what you need.",
        "key_concepts": [
            ("Table", "A drawer in the filing cabinet. One table for 'Students', another for 'Courses'."),
            ("Row/Record", "A single folder in the drawer — one complete set of information about one thing."),
            ("Column/Field", "A label on the folder — 'Name', 'Age', 'Grade'. Every folder has the same labels."),
            ("Primary Key", "A unique ID like a Social Security Number. No two people can have the same one."),
            ("Foreign Key", "A reference to another drawer. Like a folder saying 'See drawer 3, folder 7 for details'."),
            ("Query", "A specific question you ask: 'Show me all students who scored above 90 in Math'."),
        ],
        "real_world_examples": [
            ("Library Card Catalog", "Old libraries had drawers (Tables) with cards (Rows). Each card had fields: Title, Author, ISBN. The ISBN is the Primary Key — unique for every book."),
            ("Excel Spreadsheet", "Columns A, B, C are fields. Each row is a record. But a database is smarter — it can link multiple sheets and enforce rules."),
            ("Amazon Orders", "When you order, Amazon creates a row in 'Orders' table. It uses Foreign Keys to link to your 'Customer' profile and the 'Product' catalog without duplicating data."),
            ("Hospital Records", "Your patient ID (Primary Key) links your 'Appointments', 'Prescriptions', and 'Lab Results' tables. One change to your address updates everywhere automatically."),
        ],
        "pros_cons": [
            ("Data Integrity", "Rules prevent bad data. Can't delete a customer who still has orders."),
            ("Powerful Queries", "Ask complex questions across millions of rows in milliseconds."),
            ("Learning Curve", "SQL syntax and database design require practice."),
            ("Setup Overhead", "More complex than a simple text file for tiny projects."),
        ],
        "code_example": """-- Create a table (drawer)
CREATE TABLE Students (
    student_id INTEGER PRIMARY KEY,  -- Unique ID
    name TEXT NOT NULL,
    age INTEGER,
    grade TEXT
);

-- Insert data (add folders)
INSERT INTO Students VALUES 
    (1, 'Alice', 20, 'A'),
    (2, 'Bob', 21, 'B'),
    (3, 'Carol', 19, 'A');

-- Query: Find A-grade students
SELECT name, age FROM Students 
WHERE grade = 'A';
-- Result: Alice, 20 | Carol, 19

-- Query: Average age
SELECT AVG(age) FROM Students;
-- Result: 20.0

-- Update Carol's age
UPDATE Students 
SET age = 20 
WHERE name = 'Carol';""",
        "diagram_type": "database_schema"
    },

    "sorting": {
        "title": "Sorting Algorithms",
        "overview": "Sorting is like organizing a hand of playing cards. You pick up cards one by one and place them in the right position so you can quickly find what you need. Different strategies work better depending on how messy your hand is!",
        "key_concepts": [
            ("Comparison-Based", "Most sorts work by comparing two elements: 'Is A bigger than B?'"),
            ("Stable Sort", "If two items are equal, their original order is preserved. Like keeping duplicate cards in the order you received them."),
            ("In-Place", "Sorting without extra space — rearranging cards on the table instead of using a second table."),
            ("Time Complexity", "How long it takes: O(n^2) for small/simple sorts, O(n log n) for efficient ones."),
        ],
        "real_world_examples": [
            ("Playing Cards", "Insertion Sort: You pick up cards one by one and insert each into its correct position in your hand."),
            ("Library Books", "Merge Sort: Divide books into small piles, sort each pile, then merge sorted piles together. Librarians use this for massive collections!"),
            ("Sports Rankings", "Quick Sort: Pick a team as 'pivot'. Put all worse teams on the left, better on the right. Repeat for each side. Used in live tournament brackets."),
            ("Gradebook", "Bubble Sort: Compare adjacent students' grades and swap if out of order. Repeat until no swaps needed. Simple but slow for big classes."),
        ],
        "pros_cons": [
            ("Essential Skill", "Searching is faster on sorted data. Databases, file systems, and APIs all rely on sorting."),
            ("Algorithm Selection Matters", "Quicksort for speed, Mergesort for stability, Counting sort for integers."),
            ("O(n^2) Can Be Slow", "Bubble/Insertion sort struggle with 1 million items."),
            ("Not Always Needed", "Python's built-in sort() is highly optimized — rarely need to write your own."),
        ],
        "code_example": """# Bubble Sort - Simple but slow
 def bubble_sort(arr):
     n = len(arr)
     for i in range(n):
         for j in range(0, n - i - 1):
             if arr[j] > arr[j + 1]:
                 arr[j], arr[j + 1] = arr[j + 1], arr[j]
     return arr

 # Quick Sort - Fast divide & conquer
 def quick_sort(arr):
     if len(arr) <= 1:
         return arr
     pivot = arr[len(arr) // 2]
     left = [x for x in arr if x < pivot]
     middle = [x for x in arr if x == pivot]
     right = [x for x in arr if x > pivot]
     return quick_sort(left) + middle + quick_sort(right)

 scores = [64, 34, 25, 12, 22, 11, 90]
 print(quick_sort(scores))  # [11, 12, 22, 25, 34, 64, 90]""",
        "diagram_type": "sorting_comparison"
    },

    "hash table": {
        "title": "Hash Tables & Hashing",
        "overview": "A Hash Table is like a giant wall of numbered mailboxes at an apartment building. Instead of searching every mailbox to find yours, the building manager uses your name (key) to calculate your mailbox number (hash). You go directly to that mailbox — instant access!",
        "key_concepts": [
            ("Key", "Your name or ID — the thing you want to store/lookup."),
            ("Hash Function", "The calculation that converts your name into a mailbox number. Good hash functions spread people evenly across mailboxes."),
            ("Bucket/Slot", "The actual mailbox where the value is stored."),
            ("Collision", "When two names hash to the same mailbox number. Like two Smiths wanting mailbox #42. Solved by 'chaining' (linked list in the mailbox) or 'open addressing' (find next empty mailbox)."),
            ("Load Factor", "How full the mailboxes are. Above 70% full, collisions become frequent and performance drops."),
        ],
        "real_world_examples": [
            ("Apartment Mailboxes", "Manager uses a formula on your last name to assign a box. You don't search 500 boxes — you know exactly where to go."),
            ("Python Dictionaries", "Every dict in Python is a hash table. my_dict['name'] hashes 'name' to find the value instantly — O(1) time."),
            ("Password Storage", "Websites don't store your password. They store its hash. When you log in, they hash your input and compare. Even if hackers steal the database, they can't reverse the hash."),
            ("Library ISBN Lookup", "Scan an ISBN barcode (key), the computer hashes it to find the book's location (value) in the database instantly."),
        ],
        "pros_cons": [
            ("Lightning Fast", "O(1) average for insert, delete, and lookup. Direct mailbox access."),
            ("Flexible Keys", "Can use strings, tuples, or custom objects as keys."),
            ("Collision Handling Needed", "Poor hash functions cause clustering — many people fighting for the same mailbox."),
            ("Unordered", "Items aren't stored in any meaningful order. No 'get the 5th item' without extra work."),
        ],
        "code_example": """# Python dict IS a hash table!
student_grades = {
    "Alice": 95,
    "Bob": 87,
    "Carol": 92
}

# O(1) lookup - direct mailbox access!
print(student_grades["Alice"])  # 95

# O(1) insertion
student_grades["David"] = 88

# Check existence - O(1)
if "Bob" in student_grades:
    print("Bob is enrolled")

# Iterating (order not guaranteed in older Python)
for name, grade in student_grades.items():
    print(f"{name}: {grade}")

# Custom hash table concept
class SimpleHashTable:
    def __init__(self, size=10):
        self.size = size
        self.buckets = [[] for _ in range(size)]

    def _hash(self, key):
        return hash(key) % self.size

    def put(self, key, value):
        index = self._hash(key)
        for i, (k, v) in enumerate(self.buckets[index]):
            if k == key:
                self.buckets[index][i] = (key, value)
                return
        self.buckets[index].append((key, value))

    def get(self, key):
        index = self._hash(key)
        for k, v in self.buckets[index]:
            if k == key:
                return v
        return None""",
        "diagram_type": "hash_table_structure"
    },

    "graph": {
        "title": "Graph Data Structure",
        "overview": "A Graph is like a map of cities connected by roads. Cities are 'nodes' (or vertices), and roads are 'edges'. Some roads are one-way (directed), others are two-way (undirected). The GPS uses graphs to find the shortest route from your home to the restaurant!",
        "key_concepts": [
            ("Vertex/Node", "A city on the map — a point in the graph."),
            ("Edge", "A road connecting two cities. Can have a weight (distance in km) or be unweighted."),
            ("Directed Graph", "One-way streets. Edge A->B exists, but B->A might not. Like a one-way road."),
            ("Undirected Graph", "Two-way streets. If A connects to B, then B automatically connects to A."),
            ("Adjacency List", "For each city, keep a list of directly connected cities. Efficient for sparse maps."),
            ("Adjacency Matrix", "A grid where cell [i][j] = 1 if road exists between city i and j. Better for dense maps."),
        ],
        "real_world_examples": [
            ("Google Maps", "Cities = nodes, Roads = edges with weights (distance/time). Dijkstra's algorithm finds the shortest path — that's how GPS works!"),
            ("Social Networks", "People are nodes. A friendship is an edge. Facebook suggests friends using 'mutual friends' — nodes connected through common neighbors."),
            ("Flight Routes", "Airports are nodes, flights are edges. Airlines use graph algorithms to optimize routes and minimize fuel costs."),
            ("Internet Routing", "Routers are nodes, connections are edges. BGP protocol finds the best path for your data packets through the internet graph."),
        ],
        "pros_cons": [
            ("Models Real World", "Perfect for networks, maps, and relationships."),
            ("Powerful Algorithms", "Shortest path, minimum spanning tree, cycle detection solve countless problems."),
            ("Complex Implementation", "More complex than trees or lists."),
            ("Memory Intensive", "Dense graphs need O(V^2) space with adjacency matrices."),
        ],
        "code_example": """from collections import defaultdict, deque

class Graph:
    def __init__(self):
        self.adjacency_list = defaultdict(list)

    def add_edge(self, u, v, directed=False):
        self.adjacency_list[u].append(v)
        if not directed:
            self.adjacency_list[v].append(u)

    def bfs(self, start):
        # Breadth-First Search: Explore all neighbors first
        visited = set([start])
        queue = deque([start])
        result = []

        while queue:
            node = queue.popleft()
            result.append(node)
            for neighbor in self.adjacency_list[node]:
                if neighbor not in visited:
                    visited.add(neighbor)
                    queue.append(neighbor)
        return result

    def dfs(self, start, visited=None):
        # Depth-First Search: Go deep before going wide
        if visited is None:
            visited = set()
        visited.add(start)
        result = [start]

        for neighbor in self.adjacency_list[start]:
            if neighbor not in visited:
                result.extend(self.dfs(neighbor, visited))
        return result

# Create a social network graph
social = Graph()
social.add_edge("Alice", "Bob")
social.add_edge("Alice", "Carol")
social.add_edge("Bob", "David")

print("BFS from Alice:", social.bfs("Alice"))
# Alice -> Bob, Carol -> David

print("DFS from Alice:", social.dfs("Alice"))
# Alice -> Bob -> David -> Carol (or similar deep path)""",
        "diagram_type": "graph_structure"
    },

    "tree": {
        "title": "Trees & Binary Trees",
        "overview": "A Tree is like a family tree or an organization chart. It starts with one boss (root) at the top. Each person can have subordinates (children), but everyone has exactly one boss (parent) — except the CEO. There are no loops; you can't be your own grandparent!",
        "key_concepts": [
            ("Root", "The CEO — the topmost node with no parent. Every tree has exactly one root."),
            ("Node", "Any person in the org chart — contains data and links to subordinates."),
            ("Parent & Child", "Your boss is your parent; your direct reports are your children."),
            ("Leaf", "An employee with no subordinates — a node with no children."),
            ("Binary Tree", "Each person has at most 2 direct reports (left child and right child)."),
            ("Binary Search Tree (BST)", "A special binary tree where left subordinates are 'less than' the boss, and right subordinates are 'greater than'. Makes searching super fast!"),
            ("Height/Depth", "How many levels from root to the deepest leaf. A flat org chart has low height; a tall one has high height."),
        ],
        "real_world_examples": [
            ("Company Org Chart", "CEO at root, VPs as children, Managers as grandchildren. Shows hierarchy and reporting structure clearly."),
            ("File System", "C: drive is root. Folders are internal nodes. Files are leaves. 'C:\\Users\\Alice\\Documents\\resume.pdf' is a path from root to leaf."),
            ("HTML DOM", "Web pages are trees! <html> is root, <body> is child, <div>s and <p>s are descendants. JavaScript traverses this tree to modify the page."),
            ("Auto-Complete", "When you type in Google, a Trie (special tree) quickly suggests completions by traversing paths through stored words."),
        ],
        "pros_cons": [
            ("Hierarchical Structure", "Naturally models real-world hierarchies — org charts, file systems, categories."),
            ("Fast Search (BST)", "O(log n) search in balanced BSTs — like binary search but with dynamic insertion."),
            ("Can Become Unbalanced", "If you insert sorted data (1, 2, 3, 4, 5), a BST becomes a linked list — O(n) search."),
            ("No Random Access", "Must traverse from root; no direct index access like arrays."),
        ],
        "code_example": """class TreeNode:
    def __init__(self, value):
        self.value = value
        self.left = None
        self.right = None

class BinarySearchTree:
    def __init__(self):
        self.root = None

    def insert(self, value):
        if not self.root:
            self.root = TreeNode(value)
            return
        self._insert_recursive(self.root, value)

    def _insert_recursive(self, node, value):
        if value < node.value:
            if node.left is None:
                node.left = TreeNode(value)
            else:
                self._insert_recursive(node.left, value)
        else:
            if node.right is None:
                node.right = TreeNode(value)
            else:
                self._insert_recursive(node.right, value)

    def search(self, value):
        return self._search_recursive(self.root, value)

    def _search_recursive(self, node, value):
        if node is None or node.value == value:
            return node
        if value < node.value:
            return self._search_recursive(node.left, value)
        return self._search_recursive(node.right, value)

    def inorder(self, node=None, result=None):
        if result is None:
            result = []
            node = self.root
        if node:
            self.inorder(node.left, result)
            result.append(node.value)
            self.inorder(node.right, result)
        return result

# Build a BST
bst = BinarySearchTree()
for val in [50, 30, 70, 20, 40, 60, 80]:
    bst.insert(val)

print(bst.inorder())  # [20, 30, 40, 50, 60, 70, 80] - sorted!
print(bst.search(40).value if bst.search(40) else "Not found")  # 40""",
        "diagram_type": "tree_structure"
    },

    "time complexity": {
        "title": "Time Complexity & Big O",
        "overview": "Time Complexity is like comparing different routes to school. Big O notation tells you how the travel time grows as the school gets farther away — not the exact minutes, but the pattern. O(1) means constant time (same effort regardless of distance), while O(n^2) means effort explodes as distance increases!",
        "key_concepts": [
            ("Big O Notation", "Describes the upper bound of growth. 'In the worst case, how bad can it get?'"),
            ("O(1) - Constant", "Same time regardless of input size. Like grabbing the first book from a shelf — whether there are 10 or 10 million books."),
            ("O(log n) - Logarithmic", "Doubling the input barely increases time. Like binary search — finding one word in a 1000-page vs 2000-page dictionary takes just 1 more step."),
            ("O(n) - Linear", "Time grows proportionally. Checking every student in a class of 30 takes 30 checks; a class of 60 takes 60."),
            ("O(n log n) - Linearithmic", "Slightly worse than linear. Efficient sorting algorithms like Merge Sort and Quick Sort live here."),
            ("O(n^2) - Quadratic", "Time explodes! Nested loops. Checking every pair of students in a class: 30 students = 900 pairs, 60 students = 3600 pairs."),
            ("O(2^n) - Exponential", "Nightmare territory. Adding ONE item doubles the time. Used only when absolutely necessary."),
        ],
        "real_world_examples": [
            ("Pizza Party", "O(1): Grabbing a slice from the box (instant). O(n): Handing out slices one by one to n people. O(n^2): Every person shakes hands with every other person."),
            ("Phone Book Search", "O(n): Checking every name from page 1. O(log n): Binary search — open middle, eliminate half. O(1): Using the index tab (if it exists)."),
            ("Password Cracking", "O(2^n): Trying every possible combination. Each extra character in the password makes it exponentially harder to crack!"),
            ("Social Media Friends", "O(n): Counting your friends. O(n^2): Suggesting 'People You May Know' by checking mutual friends between every pair."),
        ],
        "pros_cons": [
            ("Universal Language", "Big O lets engineers compare algorithms without running code."),
            ("Focus on Growth", "Shows what happens at scale — crucial for millions of users."),
            ("Hides Constants", "O(100n) and O(n) are both O(n), but the first is 100x slower in practice."),
            ("Worst-Case Focus", "Sometimes average case is much better than worst case."),
        ],
        "code_example": """# O(1) - Constant time
def get_first_item(arr):
    return arr[0]  # Instant, regardless of array size

# O(n) - Linear time
def find_max(arr):
    maximum = arr[0]
    for num in arr:  # Check each once
        if num > maximum:
            maximum = num
    return maximum

# O(n^2) - Quadratic time
def find_duplicate_pairs(arr):
    pairs = []
    for i in range(len(arr)):      # n times
        for j in range(i+1, len(arr)):  # n times
            if arr[i] == arr[j]:
                pairs.append((arr[i], arr[j]))
    return pairs

# O(log n) - Logarithmic
def binary_search(arr, target):
    left, right = 0, len(arr) - 1
    while left <= right:
        mid = (left + right) // 2
        if arr[mid] == target:
            return mid
        elif arr[mid] < target:
            left = mid + 1
        else:
            right = mid - 1
    return -1""",
        "diagram_type": "complexity_chart"
    },
}

# ============================================================
# AI CONTENT GENERATION (Optional)
# ============================================================

def generate_with_ai(topic: str) -> Dict:
    """Generate structured notes using OpenAI or Anthropic API."""

    system_prompt = """You are an expert Computer Science educator who creates beautiful, beginner-friendly study notes.

    Generate structured notes for the given CS topic. Return ONLY a JSON object with this exact structure:
    {
        "title": "Topic Name",
        "overview": "A simple, engaging 3-4 sentence explanation using a real-world analogy. Use very simple language.",
        "key_concepts": [
            ["Concept Name", "Simple explanation with a real-world analogy. Keep it under 3 sentences."],
            ...
        ],
        "real_world_examples": [
            ["Example Title", "Detailed real-world scenario explaining how this concept applies. Use everyday situations. 2-3 sentences."],
            ...
        ],
        "pros_cons": [
            ["Advantage/Disadvantage", "Explanation with real-world context."],
            ...
        ],
        "code_example": "A clean, well-commented code example in Python. Include comments explaining each part.",
        "diagram_type": "One of: array_structure, linked_list, stack_operations, queue_structure, binary_search_tree, recursion_tree, oop_hierarchy, database_schema, sorting_comparison, hash_table_structure, graph_structure, tree_structure, complexity_chart"
    }

    Rules:
    - Language must be extremely simple. A 12-year-old should understand.
    - Every concept MUST have a real-world analogy.
    - Include 4-6 key concepts, 3-4 real-world examples, 4-5 pros/cons.
    - Code should be practical and runnable.
    - The overview should hook the reader with an interesting analogy."""

    user_prompt = f"Generate comprehensive study notes for the Computer Science topic: '{topic}'"

    try:
        if OPENAI_API_KEY and REQUESTS_AVAILABLE:
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": 0.7,
                    "response_format": {"type": "json_object"}
                },
                timeout=30
            )
            data = response.json()
            content = data["choices"][0]["message"]["content"]
            return json.loads(content)

        elif ANTHROPIC_API_KEY and REQUESTS_AVAILABLE:
            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "Content-Type": "application/json",
                    "anthropic-version": "2023-06-01"
                },
                json={
                    "model": "claude-3-haiku-20240307",
                    "max_tokens": 4000,
                    "system": system_prompt,
                    "messages": [{"role": "user", "content": user_prompt}]
                },
                timeout=30
            )
            data = response.json()
            content = data["content"][0]["text"]
            start = content.find("{")
            end = content.rfind("}") + 1
            return json.loads(content[start:end])

    except Exception as e:
        print(f"AI generation failed: {e}")
        return None

    return None


def get_topic_data(topic: str) -> Dict:
    """Get topic data from knowledge base or AI."""
    topic_lower = topic.lower().strip()

    # Check knowledge base
    for key, data in KNOWLEDGE_BASE.items():
        if key in topic_lower or topic_lower in key:
            print(f"Found '{key}' in knowledge base!")
            return data

    # Try AI if available
    if USE_AI and REQUESTS_AVAILABLE:
        print("Generating with AI...")
        ai_data = generate_with_ai(topic)
        if ai_data:
            return ai_data

    # Fallback: create generic template
    print("Topic not in knowledge base and no AI key found. Using generic template.")
    return {
        "title": topic.title(),
        "overview": f"{topic.title()} is a fundamental concept in Computer Science. Understanding it well opens doors to building better software and solving complex problems efficiently.",
        "key_concepts": [
            ("Definition", f"The formal definition and core idea behind {topic.title()}."),
            ("Importance", f"Why {topic.title()} matters in real-world software development."),
            ("Implementation", f"How {topic.title()} is typically implemented in code."),
            ("Applications", f"Common scenarios where {topic.title()} is used."),
        ],
        "real_world_examples": [
            ("Software Development", f"Real applications of {topic.title()} in modern software."),
            ("Daily Life", f"How {topic.title()} concepts appear in everyday technology."),
        ],
        "pros_cons": [
            ("Efficiency", "Can improve performance when applied correctly."),
            ("Versatility", "Applicable across many domains and languages."),
            ("Learning Curve", "May require time to master fully."),
            ("Context Dependent", "Not always the best solution for every problem."),
        ],
        "code_example": f"# Example code for {topic}\n# Please research and add specific implementation\nprint('Study {topic} thoroughly!')",
        "diagram_type": "array_structure"
    }


# ============================================================
# DIAGRAM GENERATION
# ============================================================

def create_modern_diagram(diagram_type: str, topic: str, output_path: str) -> str:
    """Generate a modern-styled diagram and save to output_path. Returns path or None."""
    if not MATPLOTLIB_AVAILABLE:
        return None

    plt.style.use('seaborn-v0_8-whitegrid')
    fig, ax = plt.subplots(figsize=(10, 6), facecolor='#F8F9FA')
    ax.set_facecolor('#F8F9FA')
    ax.axis('off')

    color_primary = '#2D5AF5'
    color_secondary = '#7B61FF'
    color_accent = '#00C9A7'
    color_dark = '#1A1A2E'
    color_light = '#E8EAF6'
    color_muted = '#636E72'
    color_warning = '#FF6B6B'

    if diagram_type == "array_structure":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Array Structure', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        boxes = ['Index 0', 'Index 1', 'Index 2', 'Index 3', 'Index 4']
        values = ['A', 'B', 'C', 'D', 'E']
        colors = [color_primary, color_secondary, color_accent, color_primary, color_secondary]

        for i, (box, val, col) in enumerate(zip(boxes, values, colors)):
            x = 1 + i * 1.6
            rect = FancyBboxPatch((x, 2.5), 1.4, 1.2, boxstyle="round,pad=0.1", 
                                   facecolor=col, edgecolor='white', linewidth=2, alpha=0.9)
            ax.add_patch(rect)
            ax.text(x + 0.7, 3.1, val, fontsize=14, fontweight='bold', ha='center', color='white')
            ax.text(x + 0.7, 2.2, box, fontsize=9, ha='center', color=color_dark, alpha=0.7)

        ax.annotate('O(1) Access', xy=(5, 1.5), fontsize=11, ha='center', 
                   color=color_dark, style='italic')
        ax.annotate('Contiguous Memory', xy=(5, 1.0), fontsize=10, ha='center', 
                   color=color_muted)

    elif diagram_type == "linked_list":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Linked List Structure', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        nodes = [('A', color_primary), ('B', color_secondary), ('C', color_accent), ('D', color_primary)]
        for i, (val, col) in enumerate(nodes):
            x = 1 + i * 2.2
            rect1 = FancyBboxPatch((x, 2.5), 1, 1.2, boxstyle="round,pad=0.05", 
                                    facecolor=col, edgecolor='white', linewidth=2, alpha=0.9)
            ax.add_patch(rect1)
            ax.text(x + 0.5, 3.1, val, fontsize=14, fontweight='bold', ha='center', color='white')

            rect2 = FancyBboxPatch((x + 1, 2.5), 0.6, 1.2, boxstyle="round,pad=0.05", 
                                    facecolor=color_light, edgecolor=col, linewidth=2)
            ax.add_patch(rect2)
            ax.text(x + 1.3, 3.1, '->', fontsize=16, ha='center', color=col)

            if i < len(nodes) - 1:
                ax.annotate('', xy=(x + 2.2, 3.1), xytext=(x + 1.6, 3.1),
                           arrowprops=dict(arrowstyle='->', color=col, lw=2))

        ax.text(8.8, 3.1, 'NULL', fontsize=12, ha='center', color=color_muted, style='italic')
        ax.annotate('Head', xy=(1.5, 4.2), fontsize=10, ha='center', color=color_dark,
                   arrowprops=dict(arrowstyle='->', color=color_primary, lw=1.5))

    elif diagram_type == "stack_operations":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Stack (LIFO)', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        stack_items = ['Pop <-', 'C', 'B', 'A (Bottom)']
        colors_stack = [color_warning, color_accent, color_secondary, color_primary]

        for i, (item, col) in enumerate(zip(stack_items, colors_stack)):
            y = 1 + i * 0.9
            rect = FancyBboxPatch((3.5, y), 3, 0.7, boxstyle="round,pad=0.05", 
                                   facecolor=col, edgecolor='white', linewidth=2, alpha=0.9)
            ax.add_patch(rect)
            ax.text(5, y + 0.35, item, fontsize=12, fontweight='bold', ha='center', color='white')

        ax.annotate('Push (Add)', xy=(7.5, 4.2), fontsize=11, color=color_dark,
                   arrowprops=dict(arrowstyle='->', color=color_accent, lw=2))
        ax.annotate('Pop (Remove)', xy=(7.5, 3.5), fontsize=11, color=color_dark,
                   arrowprops=dict(arrowstyle='->', color=color_warning, lw=2))
        ax.text(5, 0.5, 'Last In, First Out — Like a stack of plates!', 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "queue_structure":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Queue (FIFO)', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        queue_items = ['A', 'B', 'C', 'D', 'E']
        colors_q = [color_primary, color_secondary, color_accent, color_secondary, color_primary]

        for i, (item, col) in enumerate(zip(queue_items, colors_q)):
            x = 1.5 + i * 1.4
            rect = FancyBboxPatch((x, 2.5), 1.2, 1, boxstyle="round,pad=0.05", 
                                   facecolor=col, edgecolor='white', linewidth=2, alpha=0.9)
            ax.add_patch(rect)
            ax.text(x + 0.6, 3.0, item, fontsize=14, fontweight='bold', ha='center', color='white')

        ax.annotate('Dequeue <-', xy=(0.8, 3.0), fontsize=11, color=color_warning,
                   arrowprops=dict(arrowstyle='->', color=color_warning, lw=2))
        ax.annotate('-> Enqueue', xy=(8.5, 3.0), fontsize=11, color=color_accent,
                   arrowprops=dict(arrowstyle='->', color=color_accent, lw=2))
        ax.text(5, 1.5, 'First In, First Out — Like a line at the movies!', 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "binary_search_tree":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Binary Search', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        positions = [(5, 4.5, '50'), (3, 3.2, '30'), (7, 3.2, '70'), (2, 2, '20'), (4, 2, '40'), (6, 2, '60'), (8, 2, '80')]
        for x, y, val in positions:
            circle = plt.Circle((x, y), 0.4, facecolor=color_primary if int(val) == 50 else color_secondary, 
                               edgecolor='white', linewidth=2)
            ax.add_patch(circle)
            ax.text(x, y, val, fontsize=11, fontweight='bold', ha='center', va='center', color='white')

        connections = [((5, 4.1), (3, 3.6)), ((5, 4.1), (7, 3.6)), ((3, 2.8), (2, 2.4)), 
                      ((3, 2.8), (4, 2.4)), ((7, 2.8), (6, 2.4)), ((7, 2.8), (8, 2.4))]
        for start, end in connections:
            ax.plot([start[0], end[0]], [start[1], end[1]], color=color_muted, linewidth=2, zorder=0)

        ax.text(5, 1.0, 'Left < Parent < Right — Sorted for fast search!', 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "recursion_tree":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Recursion Tree', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        levels = [
            [(5, 4.5, 'f(5)')],
            [(3.5, 3.2, 'f(4)'), (6.5, 3.2, 'f(4)')],
            [(2.5, 2, 'f(3)'), (4.5, 2, 'f(3)'), (5.5, 2, 'f(3)'), (7.5, 2, 'f(3)')],
        ]

        for level in levels:
            for x, y, val in level:
                rect = FancyBboxPatch((x-0.4, y-0.25), 0.8, 0.5, boxstyle="round,pad=0.05", 
                                       facecolor=color_primary, edgecolor='white', linewidth=2)
                ax.add_patch(rect)
                ax.text(x, y, val, fontsize=9, fontweight='bold', ha='center', va='center', color='white')

        ax.annotate('', xy=(3.5, 3.5), xytext=(4.6, 4.2),
                   arrowprops=dict(arrowstyle='->', color=color_secondary, lw=1.5))
        ax.annotate('', xy=(6.5, 3.5), xytext=(5.4, 4.2),
                   arrowprops=dict(arrowstyle='->', color=color_secondary, lw=1.5))
        ax.text(5, 1.0, 'Each call breaks into smaller sub-problems until base case!', 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "oop_hierarchy":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — OOP Inheritance', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        rect_parent = FancyBboxPatch((3.5, 4), 3, 0.8, boxstyle="round,pad=0.1", 
                                      facecolor=color_primary, edgecolor='white', linewidth=2)
        ax.add_patch(rect_parent)
        ax.text(5, 4.4, 'Vehicle (Parent Class)', fontsize=12, fontweight='bold', ha='center', color='white')

        children = [('Car', 1.5, color_secondary), ('Motorcycle', 5, color_accent), ('Boat', 8.5, color_secondary)]
        for name, x, col in children:
            rect = FancyBboxPatch((x-0.8, 2.2), 1.6, 0.7, boxstyle="round,pad=0.1", 
                                   facecolor=col, edgecolor='white', linewidth=2)
            ax.add_patch(rect)
            ax.text(x, 2.55, name, fontsize=11, fontweight='bold', ha='center', color='white')
            ax.plot([5, x], [4, 2.9], color=color_muted, linewidth=2, zorder=0)

        ax.text(5, 1.2, 'Child classes inherit from Parent — reuse + extend!', 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "database_schema":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Database Schema', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        tables = [
            ('Students', 2, 3.5, color_primary, ['student_id (PK)', 'name', 'age']),
            ('Courses', 8, 3.5, color_secondary, ['course_id (PK)', 'title', 'credits']),
            ('Enrollments', 5, 1.5, color_accent, ['student_id (FK)', 'course_id (FK)', 'grade']),
        ]

        for name, x, y, col, cols in tables:
            rect = FancyBboxPatch((x-1.2, y-0.6), 2.4, 1.2, boxstyle="round,pad=0.1", 
                                   facecolor=col, edgecolor='white', linewidth=2, alpha=0.9)
            ax.add_patch(rect)
            ax.text(x, y + 0.3, name, fontsize=11, fontweight='bold', ha='center', color='white')
            ax.text(x, y - 0.1, ' | '.join(cols), fontsize=7, ha='center', color='white', alpha=0.9)

        ax.plot([2, 5], [3.5, 2.1], color=color_muted, linewidth=1.5, linestyle='--', zorder=0)
        ax.plot([8, 5], [3.5, 2.1], color=color_muted, linewidth=1.5, linestyle='--', zorder=0)
        ax.text(5, 0.8, 'Tables linked by Foreign Keys (FK) — no data duplication!', 
               fontsize=10, ha='center', color=color_muted, style='italic')

    elif diagram_type == "sorting_comparison":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Sorting Comparison', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        algorithms = [('Bubble', 'O(n^2)', 1.5, color_warning), 
                      ('Insertion', 'O(n^2)', 3.5, color_warning),
                      ('Merge', 'O(n log n)', 5.5, color_accent),
                      ('Quick', 'O(n log n)', 7.5, color_accent)]

        for name, complexity, x, col in algorithms:
            height = 1.5 if 'log' in complexity else 3.5
            rect = FancyBboxPatch((x-0.6, 1.5), 1.2, height, boxstyle="round,pad=0.05", 
                                   facecolor=col, edgecolor='white', linewidth=2, alpha=0.9)
            ax.add_patch(rect)
            ax.text(x, 1.5 + height/2, name, fontsize=10, fontweight='bold', ha='center', color='white')
            ax.text(x, 1.2, complexity, fontsize=9, ha='center', color=color_dark)

        ax.text(5, 0.7, 'Faster algorithms (green) scale much better with large data!', 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "hash_table_structure":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Hash Table', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        ax.text(1.5, 4.5, "Key: 'Alice'", fontsize=10, ha='center', color=color_dark,
               bbox=dict(boxstyle='round', facecolor=color_light, edgecolor=color_primary))
        ax.annotate('', xy=(3.5, 4.5), xytext=(2.5, 4.5),
                   arrowprops=dict(arrowstyle='->', color=color_primary, lw=2))
        ax.text(4.5, 4.5, "Hash Function", fontsize=10, ha='center', color=color_dark,
               bbox=dict(boxstyle='round', facecolor=color_secondary, edgecolor='white', alpha=0.9))
        ax.annotate('', xy=(6.5, 4.5), xytext=(5.5, 4.5),
                   arrowprops=dict(arrowstyle='->', color=color_secondary, lw=2))

        buckets = [('0', color_muted), ('1', color_muted), ('2', color_accent), ('3', color_muted), ('4', color_muted)]
        for i, (label, col) in enumerate(buckets):
            x = 1.5 + i * 1.6
            rect = FancyBboxPatch((x, 2.5), 1.2, 1, boxstyle="round,pad=0.05", 
                                   facecolor=col, edgecolor='white', linewidth=2, alpha=0.9)
            ax.add_patch(rect)
            ax.text(x + 0.6, 3.0, label, fontsize=12, fontweight='bold', ha='center', color='white')

        ax.text(5.3, 3.0, "95", fontsize=10, ha='center', color='white', fontweight='bold')
        ax.text(5, 1.5, "Hash('Alice') -> Bucket 2 -> Value: 95 (O(1) access!)", 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "graph_structure":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Graph Structure', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        nodes = [('A', 2, 4, color_primary), ('B', 5, 4.5, color_secondary), 
                 ('C', 8, 4, color_accent), ('D', 3, 2, color_secondary), ('E', 7, 2, color_primary)]
        for name, x, y, col in nodes:
            circle = plt.Circle((x, y), 0.4, facecolor=col, edgecolor='white', linewidth=2)
            ax.add_patch(circle)
            ax.text(x, y, name, fontsize=12, fontweight='bold', ha='center', va='center', color='white')

        edges = [('A', 'B'), ('B', 'C'), ('A', 'D'), ('D', 'E'), ('C', 'E'), ('B', 'D')]
        pos_dict = {n: (x, y) for n, x, y, _ in nodes}
        for u, v in edges:
            x1, y1 = pos_dict[u]
            x2, y2 = pos_dict[v]
            ax.plot([x1, x2], [y1, y2], color=color_muted, linewidth=2, zorder=0, alpha=0.7)

        ax.text(5, 0.8, 'Nodes connected by Edges — models real-world networks!', 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "tree_structure":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Tree Structure', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        tree_nodes = [(5, 4.5, '50', color_primary), (3, 3.2, '30', color_secondary), 
                      (7, 3.2, '70', color_secondary), (2, 2, '20', color_accent), 
                      (4, 2, '40', color_accent), (6, 2, '60', color_accent), (8, 2, '80', color_accent)]
        for x, y, val, col in tree_nodes:
            circle = plt.Circle((x, y), 0.35, facecolor=col, edgecolor='white', linewidth=2)
            ax.add_patch(circle)
            ax.text(x, y, val, fontsize=10, fontweight='bold', ha='center', va='center', color='white')

        tree_edges = [((5, 4.15), (3, 3.55)), ((5, 4.15), (7, 3.55)), 
                      ((3, 2.85), (2, 2.35)), ((3, 2.85), (4, 2.35)),
                      ((7, 2.85), (6, 2.35)), ((7, 2.85), (8, 2.35))]
        for start, end in tree_edges:
            ax.plot([start[0], end[0]], [start[1], end[1]], color=color_muted, linewidth=2, zorder=0)

        ax.text(5, 1.0, 'Root at top, children below, leaves at bottom — hierarchical data!', 
               fontsize=11, ha='center', color=color_muted, style='italic')

    elif diagram_type == "complexity_chart":
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 5.5, f'{topic} — Big O Complexity', fontsize=16, fontweight='bold', 
                ha='center', color=color_dark)

        x = np.linspace(1, 9, 50)
        ax.plot(x, 1 + np.zeros_like(x), color=color_primary, linewidth=3, label='O(1)')
        ax.plot(x, 1 + 0.3 * np.log(x), color=color_secondary, linewidth=3, label='O(log n)')
        ax.plot(x, 1 + 0.4 * x, color=color_accent, linewidth=3, label='O(n)')
        ax.plot(x, 1 + 0.08 * x**2, color=color_warning, linewidth=3, label='O(n^2)')

        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.legend(loc='upper left', frameon=True, fancybox=True, shadow=True)
        ax.set_xlabel('Input Size (n)', fontsize=11, color=color_dark)
        ax.set_ylabel('Operations', fontsize=11, color=color_dark)
        ax.text(5, 0.3, 'O(n^2) explodes quickly! Always prefer O(log n) or O(n) when possible.', 
               fontsize=10, ha='center', color=color_muted, style='italic')
        ax.grid(True, alpha=0.3)

    else:
        # Generic diagram
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 6)
        ax.text(5, 3, f'{topic}', fontsize=20, fontweight='bold', ha='center', color=color_primary)
        ax.text(5, 2.2, 'Visual diagram for this concept', fontsize=12, ha='center', color=color_muted)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='#F8F9FA')
    plt.close()
    return output_path

# ============================================================
# DOCUMENT FORMATTING HELPERS
# ============================================================

def set_cell_shading(cell, color: str):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_heading(doc, text: str, level: int = 1, color: str = COLORS["primary"]):
    """Add a modern styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(24)
        elif level == 2:
            run.font.size = Pt(18)
        else:
            run.font.size = Pt(14)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_info_box(doc, title: str, content: str, color: str = COLORS["accent"]):
    """Add a colored info box with title and content."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    cell = table.cell(0, 0)
    set_cell_shading(cell, color)

    # Title
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  {title}")
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.bold = True
    run.font.size = Pt(12)

    # Content
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(f"  {content}")
    run2.font.color.rgb = RGBColor(255, 255, 255)
    run2.font.size = Pt(10)

    # Set cell margins
    cell.width = Inches(6)
    return table


def add_code_block(doc, code: str):
    """Add a styled code block."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)
    set_cell_shading(cell, "1A1A2E")  # Dark background

    # Add code with line breaks preserved
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lines = code.strip().split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS["accent"])
        if i < len(lines) - 1:
            run = p.add_run('\n')

    cell.width = Inches(6)
    return table


def add_pros_cons_table(doc, items: List[Tuple[str, str]]):
    """Add a styled pros/cons table."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Advantage"
    hdr_cells[1].text = "Explanation"

    for cell in hdr_cells:
        set_cell_shading(cell, COLORS["primary"])
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Rows
    for title, desc in items:
        row_cells = table.add_row().cells
        row_cells[0].text = title
        row_cells[1].text = desc

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(4)

    return table


def add_concept_cards(doc, concepts: List[Tuple[str, str]]):
    """Add concept cards in a 2-column layout."""
    for i in range(0, len(concepts), 2):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for j in range(2):
            if i + j < len(concepts):
                title, desc = concepts[i + j]
                cell = table.cell(0, j)

                # Alternating colors
                colors = [COLORS["primary"], COLORS["secondary"]]
                set_cell_shading(cell, colors[j])

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"  {title}")
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(12)

                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run2 = p2.add_run(f"  {desc}")
                run2.font.color.rgb = RGBColor(255, 255, 255)
                run2.font.size = Pt(9)

                cell.width = Inches(3)
            else:
                table.cell(0, j).text = ""


def add_example_cards(doc, examples: List[Tuple[str, str]]):
    """Add real-world example cards with icons."""
    for title, desc in examples:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.cell(0, 0)
        set_cell_shading(cell, "F0F4FF")  # Light blue background

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"  Example: {title}")
        run.font.color.rgb = RGBColor.from_string(COLORS["primary"])
        run.font.bold = True
        run.font.size = Pt(12)

        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = p2.add_run(f"  {desc}")
        run2.font.color.rgb = RGBColor.from_string(COLORS["text"])
        run2.font.size = Pt(10)

        cell.width = Inches(6)


# ============================================================
# MAIN DOCUMENT GENERATOR
# ============================================================

def generate_notes_docx(topic: str, output_path: str = None) -> str:
    """Generate a beautifully formatted DOCX file for the given CS topic."""

    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    # Get topic data
    data = get_topic_data(topic)

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor.from_string(COLORS["text"])

    # Set narrow margins for modern look
    sections = doc.sections[0]
    sections.top_margin = Cm(2)
    sections.bottom_margin = Cm(2)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    # === COVER PAGE ===
    doc.add_paragraph()  # Spacer
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CS STUDY NOTES")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    run.font.bold = True

    doc.add_paragraph()

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run(data["title"])
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor.from_string(COLORS["primary"])
    run.font.bold = True

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Computer Science Engineering")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(COLORS["secondary"])

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    doc.add_page_break()

    # === OVERVIEW SECTION ===
    add_styled_heading(doc, "Overview", level=1)

    overview_para = doc.add_paragraph()
    overview_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = overview_para.add_run(data["overview"])
    run.font.size = Pt(12)
    run.italic = True

    add_info_box(doc, "Quick Tip", 
                "Read the real-world examples first if the technical terms feel overwhelming. The analogies will make everything click!",
                COLORS["accent"])

    doc.add_paragraph()  # Spacer

    # === KEY CONCEPTS SECTION ===
    add_styled_heading(doc, "Key Concepts", level=1)

    concept_para = doc.add_paragraph()
    concept_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = concept_para.add_run("These are the building blocks you need to master. Each concept includes a simple analogy to help you remember.")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_concept_cards(doc, data["key_concepts"])

    doc.add_paragraph()  # Spacer

    # === DIAGRAM SECTION ===
    add_styled_heading(doc, "Visual Diagram", level=1)

    diagram_para = doc.add_paragraph()
    diagram_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = diagram_para.add_run("A picture is worth a thousand words. Study this diagram to understand the structure at a glance.")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    # Generate and add diagram
    if MATPLOTLIB_AVAILABLE:
        diagram_path = os.path.join(tempfile.gettempdir(), f"cs_diagram_{topic.replace(' ', '_')}.png")
        create_modern_diagram(data.get("diagram_type", "array_structure"), data["title"], diagram_path)
        if os.path.exists(diagram_path):
            doc.add_picture(diagram_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        warning = doc.add_paragraph()
        warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = warning.add_run("[Install matplotlib to generate diagrams: pip install matplotlib]")
        run.font.color.rgb = RGBColor.from_string(COLORS["warning"])
        run.font.size = Pt(10)

    doc.add_paragraph()  # Spacer

    # === REAL-WORLD EXAMPLES SECTION ===
    add_styled_heading(doc, "Real-World Examples", level=1)

    example_intro = doc.add_paragraph()
    example_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = example_intro.add_run("These examples connect the theory to everyday life. If you understand these, you understand the concept!")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_example_cards(doc, data["real_world_examples"])

    doc.add_paragraph()  # Spacer

    # === PROS & CONS SECTION ===
    add_styled_heading(doc, "Pros & Cons", level=1)

    pros_para = doc.add_paragraph()
    pros_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = pros_para.add_run("No tool is perfect. Understanding when to use (and when NOT to use) this concept is crucial for good engineering.")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_pros_cons_table(doc, data["pros_cons"])

    doc.add_paragraph()  # Spacer

    # === CODE EXAMPLE SECTION ===
    add_styled_heading(doc, "Code Example", level=1)

    code_intro = doc.add_paragraph()
    code_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = code_intro.add_run("Study this working example. Try running it yourself and modify the values to see what happens!")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_code_block(doc, data["code_example"])

    doc.add_paragraph()  # Spacer

    # === SUMMARY BOX ===
    add_styled_heading(doc, "Summary", level=1)

    summary_items = [
        f"Topic: {data['title']}",
        f"Key Takeaway: {data['overview'][:100]}...",
        "Practice: Try implementing the code example from scratch without looking.",
        "Next Steps: Look for interview questions related to this topic on LeetCode or HackerRank."
    ]

    for item in summary_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(COLORS["text"])

    # === FOOTER ===
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— Generated by CS Notes Generator —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    run.italic = True

    # Save document
    if output_path is None:
        safe_topic = topic.replace(' ', '_').replace('/', '_')
        output_path = f"CS_Notes_{safe_topic}_{datetime.now().strftime('%Y%m%d')}.docx"

    doc.save(output_path)
    print(f"\nDocument saved: {os.path.abspath(output_path)}")
    return output_path


def convert_to_pdf(docx_path: str) -> str:
    """Convert DOCX to PDF using available methods."""
    pdf_path = docx_path.replace('.docx', '.pdf')

    # Try docx2pdf (Windows only)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"PDF saved: {os.path.abspath(pdf_path)}")
        return pdf_path
    except ImportError:
        pass
    except Exception as e:
        print(f"docx2pdf failed: {e}")

    # Try LibreOffice (Linux/Mac)
    try:
        import subprocess
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
             os.path.dirname(os.path.abspath(docx_path)), os.path.abspath(docx_path)],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            print(f"PDF saved: {os.path.abspath(pdf_path)}")
            return pdf_path
    except FileNotFoundError:
        pass
    except Exception as e:
        print(f"LibreOffice conversion failed: {e}")

    print("PDF conversion not available. Install docx2pdf (Windows) or LibreOffice (Linux/Mac).")
    return None


# ============================================================
# MAIN EXECUTION
# ============================================================

def main():
    print("=" * 60)
    print("   CS NOTES GENERATOR")
    print("   Modern Study Notes for Computer Science")
    print("=" * 60)
    print()

    if not DOCX_AVAILABLE:
        print("ERROR: python-docx is required.")
        print("Install with: pip install python-docx Pillow matplotlib requests")
        sys.exit(1)

    print("Enter a Computer Science topic/subject name.")
    print("Examples: Arrays, Linked List, Stack, Queue, Binary Search,")
    print("          Recursion, OOP, Database, Sorting, Hash Table,")
    print("          Graph, Tree, Time Complexity, etc.")
    print()

    topic = input("Topic: ").strip()

    if not topic:
        print("No topic entered. Exiting.")
        sys.exit(0)

    print()
    print(f"Generating notes for: {topic}")
    print("-" * 40)

    try:
        docx_path = generate_notes_docx(topic)

        print()
        print("DOCX file generated successfully!")
        print(f"Location: {os.path.abspath(docx_path)}")

        # Ask for PDF conversion
        convert = input("\nConvert to PDF? (y/n): ").strip().lower()
        if convert == 'y':
            pdf_path = convert_to_pdf(docx_path)
            if pdf_path:
                print(f"PDF Location: {os.path.abspath(pdf_path)}")

        print()
        print("Done! Open the document to view your notes.")

    except Exception as e:
        print(f"\nERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
